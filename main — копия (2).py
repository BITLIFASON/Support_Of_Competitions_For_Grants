from PyQt6 import uic
from PyQt6.QtCore import QItemSelection, QItemSelectionModel, QModelIndex
from PyQt6.QtWidgets import QApplication
from PyQt6.QtSql import QSqlDatabase, QSqlTableModel, QSqlQuery
import sys, copy
import sqlite3
from functools import partial
import itertools as it
from itertools import chain
import ast
import os
import shutil
import docx
import datetime
from datetime import datetime
# загрузка форм и окон из ui файлов
MainForm, MainWindow = uic.loadUiType("MainWindow.ui")  # главное окно
ProjViewForm, ProjViewWindow = uic.loadUiType("ProjViewWindow.ui")  # окно показа проектов
ProjSortForm, ProjSortWindow = uic.loadUiType("ProjSortWindow.ui") # окно сортировки проектов
ProjFilterForm, ProjFilterWindow = uic.loadUiType("ProjFilterWindow.ui") # окно фильтрации проектов
KonkViewForm, KonkViewWindow = uic.loadUiType("KonkViewWindow.ui")  # окно показа конкурсов
VuzViewForm, VuzViewWindow = uic.loadUiType("VuzViewWindow.ui")  # окно показа вузов
HelpViewForm, HelpViewWindow = uic.loadUiType("HelpViewWindow.ui")  # окно показа справки
HelpProgViewForm, HelpProgViewWindow = uic.loadUiType("HelpProgViewWindow.ui")  # окно показа справки по программе
DeleteWindowForm, DeleteWindow = uic.loadUiType("DeleteWindow.ui") # окно с подтверждением удаления строок
AddRowWindowForm, AddRowWindowWindow = uic.loadUiType("AddRowWindow.ui") # окно с добавлением строк
ErrorAddRowForm , ErrorAddRowWindow = uic.loadUiType("Error_row_add.ui") # окно с указанием ошибки добавлекния строки
ChangeForm , ChangeWindow = uic.loadUiType("ChangeWindow.ui") # окно с редактированием строк
AnalysisForm , AnalysisWindow = uic.loadUiType("AnalysisWindow.ui") # окно с редактированием строк
AddOkForm , ADdOkWindow = uic.loadUiType("AddOk.ui") # окно с инф о добавлении строки
FinForm , FinWindow = uic.loadUiType("FinWindow.ui") # окно с финансированием


# инициализация имени базы данных
db_name = 'bd.sqlite'


# инициализация фильтра
query_sql_filter_temp = {}
query_sql_filter_save = {}


query_sql_sort_save = {}


translate_dict_filter = {'proj.codkon':'Код конкурса', 'vuz.region':'Федеральный округ', 'vuz.oblname':'Субъект РФ', 'vuz.city':'Город', 'vuz.z2': 'Вуз'}


def add_ok():
    add_ok_window.close()


def check_active_window():
    # проверяет какое окно открыто и возвращает его номер

    for i, window in enumerate(window_list):
        if window.isActiveWindow():
            return i


def open_proj_window():
    # открывает окно показа таблицы конкурсов

    global query_sql_filter_save

    window_list[check_active_window()].close()
    proj_view_window.showMaximized()

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    proj_view_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))

    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                 proj.g1 AS "Код НИР",
                                 proj.g8 AS "Руководитель НИР",
                                 proj.g7 AS "Код по ГРНТИ",
                                 proj.z2 AS "Наименование вуза",
                                 proj.g5 AS "Плановый объем",
                                 proj.g6 AS "Наименование НИР",
                                 proj.g2 AS "Фактический объем гранта",
                                 proj.g21 as"Поквартальное финансирование",
                                 proj.g22 as"Поквартальное финансирование",
                                 proj.g23 as"Поквартальное финансирование",
                                 proj.g24 as"Поквартальное финансирование",
                                 proj.g9 as "Должность руководителя",
                                 proj.g10 as"Учёное звание руководителя",
                                 proj.g11 as"Учёная степень руководителя"
                          FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                          {cond_filter if query_sql_filter_save != {} else ""}
                          {cond_sort if query_sql_sort_save != {} else ""}""")

    # proj_view_form.tableView.setFocus()

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_view_form.tableView.setModel(table_model)
    proj_view_form.tableView.setSortingEnabled(True)
    proj_view_form.tableView.resizeColumnsToContents()


def read_data_add_row():

    global query_sql_filter_save

    # Чтение с лейблов и боксов
    table_model = QSqlTableModel()
    proj_view_form.tableView.setModel(table_model)
    data_g1 = add_view_form.lineEdit_g1.text()
    data_g7 = add_view_form.lineEdit_g7.text()
    data_g5 = add_view_form.lineEdit_g5.text()
    data_g6 = add_view_form.lineEdit_g6.toPlainText()
    data_g8 = add_view_form.lineEdit_g8.text()
    data_g9 = add_view_form.lineEdit_g9.text()
    data_g10 = add_view_form.lineEdit_g10.text()
    data_g11 = add_view_form.lineEdit_g11.text()
    data_cod_kon = add_view_form.comboBox_cod_kon.currentText()
    data_cod_vuza = add_view_form.comboBox_cod_vuz.currentText()
    j=0
    add_error_text=str()

    #Провекра вводимых данных
    # Проверка g1
    simbol="""abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\/"#$%&'()*+,-./:;<=>?@[]^_`{|}~абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ """
    if (any((c in data_g1) for c in simbol) or data_g1==''):
        j=j+1
        add_error_text=add_error_text+str('Вы ввели неверное значение g1- '+str(data_g1)+' - А ожидался код нира не из списка примечаний')
        add_view_form.lineEdit_g1.setText(str(data_g1))
        add_view_form.lineEdit_g1.setStyleSheet("color: red;")
    else:
        add_view_form.lineEdit_g1.setStyleSheet("color: black;")
    if(any((c in data_g1) for c in simbol) or data_g1!=''):
        indices = [i for i in range(0, len(comboBox_codkon1)) if comboBox_codkon1[i][0:3] == data_cod_kon[0:3]]
        proverka_g1_codkon=[]
        b = []
        for i in indices:
            b.append(comboBox_codkon1[i][3:])
        c=[]
        for i in indices:
            c.append(comboBox_codkon1[i][0:2])
        if data_g1 in b and data_cod_kon[0:2] in c:
            for i in indices:
                proverka_g1_codkon.append(comboBox_codkon1[i])
            for i in range(len(proverka_g1_codkon)):
                if int(data_g1)==int(proverka_g1_codkon[i][3:]):
                    add_error_text = add_error_text + str(
                        'Вы ввели неверное значение g1- ' + str(data_g1) + ' и  значение codkon - '+str(data_cod_kon[0:2])+' - А ожидались код НИР и код конкурса не из списка  ')
                    j=j+1
                    add_view_form.lineEdit_g1.clear()
                    add_view_form.lineEdit_g1.setText(str(data_g1))
                    add_view_form.lineEdit_g1.setStyleSheet("color: red;")
    else:
        add_view_form.lineEdit_g1.setStyleSheet("color: black;")

    stroka = [(data_g1), data_cod_kon, (data_cod_vuza),  data_g7, data_g5, data_g6, data_g8, data_g9, data_g10, data_g11]
    # проверка ввода g7
    simbol1="""abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\/"#$%&'()*+-/:;<=>?@[]^_`{|}~абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ """
    if (any((c in data_g7) for c in simbol1)):
        j=j+1
        add_error_text=add_error_text+str('\nВы ввели неверное значение g7- '+data_g7+' - А ожидалось ,например, - 50.09,50.10')
        add_view_form.lineEdit_g7.clear()
        add_view_form.lineEdit_g7.setText(str(data_g7))
        add_view_form.lineEdit_g7.setStyleSheet("color: red;")
    else:
        add_view_form.lineEdit_g7.setStyleSheet("color: black;")
    # Провекра ввода g5
    simbol2="""abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\/"#$%&'()*+-/.,:;<=>?@[]^_`{|}~абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ """
    data_g5_g24=[data_g5]
    data_g5_g24_str = ['g5']
    for i in range(len(data_g5_g24)):
        if ((any((c in data_g5_g24[i]) for c in simbol1)) or data_g5_g24[0]=='0' or data_g5_g24[0]==''):
            j=j+1
            add_error_text=add_error_text+str('\nВы ввели неверное значение '+(data_g5_g24_str[i])+' - '+str(data_g5_g24[i])+ ' - А ожидалось целочисленное значение больше 0')
            add_view_form.lineEdit_g5.clear()
            add_view_form.lineEdit_g5.setText(str(data_g5))
            add_view_form.lineEdit_g5.setStyleSheet("color: red;")
        else:
            add_view_form.lineEdit_g5.setStyleSheet("color: black;")
    #
    # проверка g6_g11
    data_g6_g11=[data_g6, data_g8, data_g9, data_g10, data_g11]
    data_g6_g11_str = ['g6', 'g8', 'g9', 'g10', 'g11']
    simbol2 = """1234567890!\/#$%&'()*+/:;<=>?@[]^_`{|}~"""
    for i in range(len(data_g6_g11)):
        if (any((c in data_g6_g11[i]) for c in simbol2)):
            j=j+1
            add_error_text = add_error_text + str(
                '\nВы ввели неверное значение ' + (data_g6_g11_str[i]) + ' - ' + str(
                    data_g6_g11[i]) + ' - А ожидалось текствое значение')
            if i == 0:
                add_view_form.lineEdit_g6.setStyleSheet("color: red;")
            if i == 1:
                add_view_form.lineEdit_g8.setStyleSheet("color: red;")
            if i == 2:
                add_view_form.lineEdit_g9.setStyleSheet("color: red;")
            if i == 3:
                add_view_form.lineEdit_g10.setStyleSheet("color: red;")
            if i == 4:
                add_view_form.lineEdit_g11.setStyleSheet("color: red;")
        else:
            if i == 0:
                add_view_form.lineEdit_g6.setStyleSheet("color: black;")
            if i == 1:
                add_view_form.lineEdit_g8.setStyleSheet("color: black;")
            if i == 2:
                add_view_form.lineEdit_g9.setStyleSheet("color: black;")
            if i == 3:
                add_view_form.lineEdit_g10.setStyleSheet("color: black;")
            if i == 4:
                add_view_form.lineEdit_g11.setStyleSheet("color: black;")

    #Вывод на экран то что значения были введены неверно
    if j>0:
        #add_view_form.comboBox_cod_kon.clear()
        #add_view_form.comboBox_cod_vuz.clear()
        #add_view_form.comboBox_cod_nir.clear()
        #add_view_form.lineEdit_g1.clear()
        #add_view_form.lineEdit_g7.clear()
        #add_view_form.lineEdit_g5.clear()
        #add_view_form.lineEdit_g6.clear()
        #add_view_form.lineEdit_g8.clear()
        #add_view_form.lineEdit_g9.clear()
        #add_view_form.lineEdit_g10.clear()
        #add_view_form.lineEdit_g11.clear()
        error_add_row_form.textBrowser.setText(add_error_text)

        error_add_row_window.show()
        #add_view_window.close()
    # Добавление элементов с таблицу
    else:

        cod_vuza=int([''.join(x) for _, x in it.groupby(data_cod_vuza, key=lambda c: c=='-')][2])
        name_vuza_add=[''.join(x) for _, x in it.groupby(data_cod_vuza, key=lambda c: c=='-')][0]
        value=[int(data_g1),data_cod_kon[0:2],cod_vuza,name_vuza_add,data_g7,data_g5,0 ,0 ,0 ,0 ,0 ,data_g6,data_g8,data_g9,data_g10,data_g11]
        for i in range(5, 11):
            if value[i] != '':
                value[i] = int(value[i])
        if value[5]==None:
            value[5]=0
        if value[5]=='':
            value[5]=0



        query = QSqlQuery()
        conn=sqlite3.connect(db_name)
        cur=conn.cursor()
        cur.execute(
            '''INSERT INTO gr_proj(g1,codkon,codvuz,z2,g7,g5,g2,g21,g22,g23,g24,g6,g8,g9,g10,g11)
             VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?);''',
            value).fetchall()
        conn.commit()
        cur.close()
        conn.close()

        table_model = QSqlTableModel()
        query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                                 proj.g1 AS "Код НИР",
                                                 proj.g8 AS "Руководитель НИР",
                                                 proj.g7 AS "Код по ГРНТИ",
                                                 proj.z2 AS "Наименование вуза",
                                                 proj.g5 AS "Плановый объем",
                                                 proj.g6 AS "Наименование НИР",
                                                 proj.g2 AS "Фактический объем гранта",
                                                 proj.g21 as"Поквартальное финансирование",
                                                 proj.g22 as"Поквартальное финансирование",
                                                 proj.g23 as"Поквартальное финансирование",
                                                 proj.g24 as"Поквартальное финансирование",
                                                 proj.g9 as "Должность руководителя",
                                                 proj.g10 as"Учёное звание руководителя",
                                                 proj.g11 as"Учёная степень руководителя"
                                          FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                          {cond_filter if query_sql_filter_save != {} else ""}
                                          {cond_sort if query_sql_sort_save != {} else ""}""")

        # ВЫДЕЛЕНИЕ СТРОКИ , КОГДА ДОБАВЛЯЕМ СТРОКУ
        table_model.setQuery(query)
        conn1 = sqlite3.connect(db_name)
        cur1 = conn1.cursor()
        cur1.execute('''select g1,codkon from gr_proj''')
        table=cur1.fetchall()
        cur1.close()
        conn1.close()
        choose_ind=int()
        for i in range(len(table)):
            if value[0] == table[i][0] and value[1]==table[i][1]:
                choose_ind=i #индекс добавляемой строки
        proj_view_window.showMaximized()

        proj_view_form.tableView.setModel(table_model)
        proj_view_form.tableView.setSortingEnabled(True)
        proj_view_form.tableView.resizeColumnsToContents()
        add_ok_window.show()

        return_back()


def add_row():

    global query_sql_filter_save

    add_view_window.show()
    proj_view_window.close()
    # заполняем какие коды уже заняты
    g2=set()
    g1=set()
    query = QSqlQuery("SELECT * FROM gr_proj")
    while query.next():
        g2.add(str(query.value(1))+'-'+str(query.value(0)))
    g2=list(g2)
    global comboBox_codkon1
    comboBox_codkon1 = []
    for j in range(1,16):
        for i in range(len(g2)):
            if int(g2[i][0:2])==j:
                comboBox_codkon1.append(g2[i])

    kod_con1 = ['01', '02', '03', '04', '05', '06', '07', '08', '09', '10', '11', '12', '13', '14', '15', '16', '17']
    comboBox_codkon1.sort()

    def spisok(a, t):
        spis3 = []
        spis2 = [[], [], [], [], [], [], [], [], [], [], [], [], [], [], [], [], []]
        spis4 = []
        spis6 = []
        spis7 = []
        spis8 = []
        k = 0
        for i in kod_con1:
            for j in range(len(a)):
                if i == a[j][0:2]:
                    spis2[k].append((a[j]))
            spis3.append(spis2[k])
            k = k + 1
        for j in range(len(spis3[t])):
            spis4.append(int(spis3[t][j][3:]))
        spis5 = sorted(spis4)
        for i in range(len(spis5)):
            spis6.append(spis4.index(spis5[i]))
        for i in spis6:
            spis7.append(spis3[t][i])
        return spis7

    indexes = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16]
    result3 = []
    for i in indexes:
        result3.append(spisok(comboBox_codkon1, i))
    result3 = list(chain.from_iterable(result3))



    def add_combo_box_change():
        add_combobox_codkon = []
        print(change_form.comboBox_cod_kon_change.currentText()[0:2])
        print(result3)
        for i in range(len(result3)):
            if result3[i][0:2] ==add_view_form.comboBox_cod_kon.currentText()[0:2]:
                add_combobox_codkon.append(result3[i])
        add_view_form.comboBox_cod_nir.clear()
        add_view_form.comboBox_cod_nir.addItems(add_combobox_codkon)
        print(add_combobox_codkon)

    add_view_form.comboBox_cod_kon.currentTextChanged.connect(add_combo_box_change)


    #заполняем код конкурса
    query = QSqlQuery("SELECT k2 , codkon FROM gr_konk")
    name_kon=set()
    while query.next():
        name_kon.add(str(query.value(1))+'-'+str(query.value(0)))
    name_kon=list(name_kon)
    add_view_form.comboBox_cod_kon.addItems(sorted(name_kon, key=lambda x:x[:3]))
    # заполняем код вуза и наименование вуза
    name_vuza=set()
    kod_vuza=set()
    query = QSqlQuery("SELECT * FROM vuz")
    while query.next():
        name_vuza.add(str(query.value(3))+ '-' + str(query.value(0)))
    name_vuza=list(name_vuza)
    add_view_form.comboBox_cod_vuz.addItems(sorted(name_vuza, key=lambda x: x[:3]))


def return_back():
    table_model = QSqlTableModel()
    proj_view_window.show()
    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                             proj.g1 AS "Код НИР",
                                             proj.g8 AS "Руководитель НИР",
                                             proj.g7 AS "Код по ГРНТИ",
                                             proj.z2 AS "Наименование вуза",
                                             proj.g5 AS "Плановый объем",
                                             proj.g6 AS "Наименование НИР",
                                             proj.g2 AS "Фактический объем гранта",
                                             proj.g21 as"Поквартальное финансирование",
                                             proj.g22 as"Поквартальное финансирование",
                                             proj.g23 as"Поквартальное финансирование",
                                             proj.g24 as"Поквартальное финансирование",
                                             proj.g9 as "Должность руководителя",
                                             proj.g10 as"Учёное звание руководителя",
                                             proj.g11 as"Учёная степень руководителя"
                                      FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                      {cond_filter if query_sql_filter_save != {} else ""}
                                      {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query)

    proj_view_form.tableView.setModel(table_model)
    proj_view_form.tableView.setSortingEnabled(True)
    proj_view_form.tableView.resizeColumnsToContents()
    add_view_form.comboBox_cod_kon.clear()
    add_view_form.comboBox_cod_vuz.clear()
    add_view_form.comboBox_cod_nir.clear()
    add_view_form.lineEdit_g1.clear()
    add_view_form.lineEdit_g7.clear()
    add_view_form.lineEdit_g5.clear()
    add_view_form.lineEdit_g6.clear()
    add_view_form.lineEdit_g8.clear()
    add_view_form.lineEdit_g9.clear()
    add_view_form.lineEdit_g10.clear()
    add_view_form.lineEdit_g11.clear()
    add_view_form.lineEdit_g1.setStyleSheet("color: black;")
    add_view_form.lineEdit_g7.setStyleSheet("color: grey;")
    add_view_form.lineEdit_g5.setStyleSheet("color: black;")
    add_view_form.lineEdit_g6.setStyleSheet("color: black;")
    add_view_form.lineEdit_g8.setStyleSheet("color: black;")
    add_view_form.lineEdit_g9.setStyleSheet("color: black;")
    add_view_form.lineEdit_g10.setStyleSheet("color: black;")
    add_view_form.lineEdit_g11.setStyleSheet("color: black;")

    add_view_window.close()
    proj_view_window.showMaximized()


def pb_no():
    delete_proj_window.close()
    proj_view_window.showMaximized()

def pb_yes():
    table_model = QSqlTableModel()
    query = QSqlQuery(f"DELETE FROM gr_proj WHERE g1=={curr_index.sibling(curr_index.row(), 1).data()} "
                      f"and gr_proj.codkon=='{curr_index.sibling(curr_index.row(), 0).data()}' ")
    table_model.setQuery(query)
    query = QSqlQuery("SELECT * FROM gr_proj")
    table_model.setQuery(query)
    delete_proj_window.close()
    proj_view_window.show()
    query = QSqlQuery("""SELECT codkon AS "Код конкурса",
                                         g1 AS "Код НИР",
                                         g8 AS "Руководитель НИР",
                                         g6 AS "Наименование НИР",
                                         g7 AS "Код по ГРНТИ",
                                         z2 AS "Наименование вуза" ,
                                         g5 AS "Плановый объем",
                                         g2 AS "Фактический объем гранта"
                                         FROM gr_proj""")

    table_model.setQuery(query)
    proj_view_form.tableView.setModel(table_model)
    proj_view_form.tableView.resizeColumnsToContents()



def proj_select_check():

    global curr_index
    global query_sql_filter_save

    table_model = QSqlTableModel()

    curr_index = proj_view_form.tableView.currentIndex()

    codkon = curr_index.sibling(curr_index.row(), 0).data()

    if curr_index.sibling(curr_index.row(), 1).data() != None:
        proj_view_window.close()
        delete_proj_window.show()

        query = QSqlQuery(f"""SELECT codkon AS "Код конкурса",
                                 g1 AS "Код НИР",
                                 g8 AS "Руководитель НИР",
                                 g7 AS "Код по ГРНТИ",
                                 z2 AS "Наименование вуза",
                                 g5 AS "Плановый объем",
                                 g6 AS "Наименование НИР",
                                 g2 AS "Фактический объем гранта",
                                 g21 as"Поквартальное финансирование",
                                 g22 as"Поквартальное финансирование",
                                 g23 as"Поквартальное финансирование",
                                 g24 as"Поквартальное финансирование",
                                 g9 as "Должность руководителя",
                                 g10 as"Учёное звание руководителя",
                                 g11 as"Учёная степень руководителя"
                             FROM gr_proj WHERE g1=={curr_index.sibling(curr_index.row(), 1).data()} and gr_proj.codkon=='{codkon}'""")
        change_form.tableView_change.resizeColumnsToContents()
        table_model.setQuery(query)
        delete_proj_form.tableView_delete_row.setModel(table_model)

    else:
        print('Ничего не выбрано')
    delete_proj_form.tableView_delete_row.resizeColumnsToContents()


def read_data_add_row1():

    global query_sql_filter_save

    curr_index1 = proj_view_form.tableView.currentIndex()
    index_change = curr_index1.row()

    table_model = QSqlTableModel()

    proj_view_form.tableView.setModel(table_model)
    data_g1 = change_form.lineEdit_g1.text()
    data_g7 = change_form.lineEdit_g7.text()
    data_g5 = change_form.lineEdit_g5.text()
    data_g2 = spis[1][6]
    data_g21 = spis[1][7]
    data_g22 = spis[1][8]
    data_g23 = spis[1][9]
    data_g24 = spis[1][10]
    data_g6 = change_form.lineEdit_g6.toPlainText()
    data_g8 = change_form.lineEdit_g8.text()
    data_g9 = change_form.lineEdit_g9.text()
    data_g10 = change_form.lineEdit_g10.text()
    data_g11 = change_form.lineEdit_g11.text()
    data_cod_kon = change_form.comboBox_cod_kon_change.currentText()
    data_cod_vuza = change_form.comboBox_cod_vuz_change.currentText()
    #
    j = 0
    add_error_text = str()
    # Провекра вводимых данных

    # Проверка g1
    simbol = """abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\/"#$%&'()*+,-./:;<=>?@[]^_`{|}~абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ """
    if (any((c in data_g1) for c in simbol) or data_g1 == ''):
        j = j + 1
        add_error_text = add_error_text + str(
            'Вы ввели неверное значение g1- ' + str(data_g1) + ' - А ожидался код НИР и код конкурса не из списка')
        change_form.lineEdit_g1.clear()
        change_form.lineEdit_g1.setText(str(data_g1))
        change_form.lineEdit_g1.setStyleSheet("color: red;")

    elif((any((c in data_g1) for c in simbol) or data_g1 != '')):
        indices = [i for i in range(0, len(comboBox_codkon)) if comboBox_codkon[i][0:3] == data_cod_kon[0:3]]
        proverka_g1_codkon=[]
        if str(res[0])==str(data_g1) and str(res[1])==str(data_cod_kon[0:2]):

            change_form.lineEdit_g1.setStyleSheet("color: black;")
        else:
            for i in indices:
                proverka_g1_codkon.append(comboBox_codkon[i])
            for i in range(len(proverka_g1_codkon)):
                if int(data_g1)==int(proverka_g1_codkon[i][3:]):
                    add_error_text = add_error_text + str(
                        'Вы ввели неверное значение g1- ' + str(data_g1) + ' и  значение codkon - '+str(data_cod_kon[0:2])+' - А ожидались код НИР и код конкурса не из списка  ')
                    j=j+1
                    change_form.lineEdit_g1.clear()
                    change_form.lineEdit_g1.setText(str(data_g1))
                    change_form.lineEdit_g1.setStyleSheet("color: red;")
    else:
        change_form.lineEdit_g1.setStyleSheet("color: black;")



    stroka = [(data_g1), data_cod_kon, (data_cod_vuza), data_g7, data_g5, (data_g2),
              (data_g21), data_g22,
              (data_g23), (data_g24), data_g6, data_g8, data_g9, data_g10, data_g11]


    # проверка ввода g7
    simbol1 = """abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\/"#$%&'()*+-/:;<=>?@[]^_`{|}~абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ """
    if (any((c in data_g7) for c in simbol1)):
        j = j + 1
        add_error_text = add_error_text + str(
            '\nВы ввели неверное значение g7- ' + data_g7 + ' - А ожидалось ,например, - 50.09.49,50.10.66')
        change_form.lineEdit_g7.clear()
        change_form.lineEdit_g7.setText(str(data_g7))
        change_form.lineEdit_g7.setStyleSheet("color: red;")
    else:
        change_form.lineEdit_g7.setStyleSheet("color: black;")
    # Провекра ввода g5-g24
    simbol2 = """abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\/#$%&'()*+/.,:;<=>?@[]^_`{|}~абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ """
    data_g5_g24 = [data_g5]
    data_g5_g24_str = ['g5']

    for i in range(len(data_g5_g24)):
        if ((any((c in data_g5_g24[i]) for c in simbol1)) or data_g5_g24[0]=='0'):
            j = j + 1
            add_error_text = add_error_text + str('\nВы ввели неверное значение ' + (data_g5_g24_str[i]) + ' - ' + str(
                data_g5_g24[i]) + ' - А ожидалось целочисленное значение больше 0')
            change_form.lineEdit_g5.clear()
            change_form.lineEdit_g5.setText(str(data_g5))
            change_form.lineEdit_g5.setStyleSheet("color: red;")
        else:
            change_form.lineEdit_g5.setStyleSheet("color: black;")

    #
    # проверка g6_g11
    data_g6_g11 = [data_g6, data_g8, data_g9, data_g10, data_g11]
    data_g6_g11_str = ['g6', 'g8', 'g9', 'g10', 'g11']
    simbol2 = """1234567890!\/#$%&'()*+/:;<=>?@[]^_`{|}~"""
    for i in range(len(data_g6_g11)):
        if (any((c in data_g6_g11[i]) for c in simbol2)):
            j = j + 1
            add_error_text = add_error_text + str(
                '\nВы ввели неверное значение ' + (data_g6_g11_str[i]) + ' - ' + str(
                    data_g6_g11[i]) + ' - А ожидалось текствое значение')
            if i ==0:
                change_form.lineEdit_g6.setStyleSheet("color: red;")
            if i==1:
                change_form.lineEdit_g8.setStyleSheet("color: red;")
            if i==2:
                change_form.lineEdit_g9.setStyleSheet("color: red;")
            if i==3:
                change_form.lineEdit_g10.setStyleSheet("color: red;")
            if i==4:
                change_form.lineEdit_g11.setStyleSheet("color: red;")
        else:
            if i ==0:
                change_form.lineEdit_g6.setStyleSheet("color: black;")
            if i==1:
                change_form.lineEdit_g8.setStyleSheet("color: black;")
            if i==2:
                change_form.lineEdit_g9.setStyleSheet("color: black;")
            if i==3:
                change_form.lineEdit_g10.setStyleSheet("color: black;")
            if i==4:
                change_form.lineEdit_g11.setStyleSheet("color: black;")





    # Вывод на экран то что значения были введены неверно
    if j > 0:
        error_add_row_form.textBrowser.setText(add_error_text)
        query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                         proj.g1 AS "Код НИР",
                                         proj.g8 AS "Руководитель НИР",
                                         proj.g7 AS "Код по ГРНТИ",
                                         proj.z2 AS "Наименование вуза",
                                         proj.g5 AS "Плановый объем",
                                         proj.g6 AS "Наименование НИР",
                                         proj.g2 AS "Фактический объем гранта",
                                         proj.g21 as"Поквартальное финансирование",
                                         proj.g22 as"Поквартальное финансирование",
                                         proj.g23 as"Поквартальное финансирование",
                                         proj.g24 as"Поквартальное финансирование",
                                         proj.g9 as "Должность руководителя",
                                         proj.g10 as"Учёное звание руководителя",
                                         proj.g11 as"Учёная степень руководителя"
                                  FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                  {cond_filter if query_sql_filter_save != {} else ""}
                                  {cond_sort if query_sql_sort_save != {} else ""}""")

        table_model.setQuery(query)

        proj_view_form.tableView.setModel(table_model)
        proj_view_form.tableView.setSortingEnabled(True)
        proj_view_form.tableView.resizeColumnsToContents()
        error_add_row_window.show()

    else:

        if data_cod_vuza != '':
            cod_vuza = int([''.join(x) for _, x in it.groupby(data_cod_vuza, key=lambda c: c == '-')][2])
            name_vuza_add = [''.join(x) for _, x in it.groupby(data_cod_vuza, key=lambda c: c == '-')][0]
        else:
            cod_vuza = ''
            name_vuza_add = ''
        value = [int(data_g1), data_cod_kon[0:2], cod_vuza, name_vuza_add, data_g7, data_g5, data_g2, data_g21,
                 data_g22, data_g23, data_g24, data_g6, data_g8, data_g9, data_g10, data_g11]

        for i in range(len(spis[1])):
            if (spis[1][i] != value[i]):
                if value[i] != '':
                    spis[1][i] = value[i]
        spis2 = spis[1]
        if data_g10=='':
            spis2[14]=None
        if data_g11=='':
            spis2[15]=None
        if data_g9=='':
            spis2[13]=None
        if data_g8=='':
            spis2[12]=None
        if data_g6=='':
            spis2[11]=None
        if data_g5=='':
            spis2[5]=None
        if data_g5=='':
            spis2[5]=0
        conn = sqlite3.connect(db_name)
        cur = conn.cursor()
        cur.execute(f'''UPDATE gr_proj SET codvuz={spis[1][2]} WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''')
        cur.execute(f'''UPDATE gr_proj SET z2=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''', (spis[1][3],))
        cur.execute(f'''UPDATE gr_proj SET g7=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''', (spis2[4],))
        cur.execute(f'''UPDATE gr_proj SET g5={spis2[5]} WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''')
        cur.execute(f'''UPDATE gr_proj SET g6=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''', (spis2[11],))
        cur.execute(f'''UPDATE gr_proj SET g8=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''', (spis2[12],))
        cur.execute(f'''UPDATE gr_proj SET g9=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''', (spis2[13],))
        cur.execute(f'''UPDATE gr_proj SET g10=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}'   ''', (spis2[14],))
        cur.execute(f'''UPDATE gr_proj SET g11=? WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}'  ''', (spis2[15],))
        cur.execute(f'''UPDATE gr_proj SET g1={spis2[0]} WHERE g1={spis[2]} and gr_proj.codkon=='{spis[3]}' ''')
        cur.execute(f'''UPDATE gr_proj SET codkon='{spis2[1]}' WHERE g1={spis2[0]} and gr_proj.codkon=='{spis[3]}' ''')
        conn.commit()
        cur.close()
        conn.close()
        change_window.close()
        proj_view_window.show()
        query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                         proj.g1 AS "Код НИР",
                                         proj.g8 AS "Руководитель НИР",
                                         proj.g7 AS "Код по ГРНТИ",
                                         proj.z2 AS "Наименование вуза",
                                         proj.g5 AS "Плановый объем",
                                         proj.g6 AS "Наименование НИР",
                                         proj.g2 AS "Фактический объем гранта",
                                         proj.g21 as"Поквартальное финансирование",
                                         proj.g22 as"Поквартальное финансирование",
                                         proj.g23 as"Поквартальное финансирование",
                                         proj.g24 as"Поквартальное финансирование",
                                         proj.g9 as "Должность руководителя",
                                         proj.g10 as"Учёное звание руководителя",
                                         proj.g11 as"Учёная степень руководителя"
                                  FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                  {cond_filter if query_sql_filter_save != {} else ""}
                                  {cond_sort if query_sql_sort_save != {} else ""}""")

        table_model.setQuery(query)

        while table_model.canFetchMore():
            table_model.fetchMore()
        proj_view_form.tableView.setFocus()
        proj_view_form.tableView.selectRow(index_change)

        proj_view_form.tableView.setModel(table_model)
        proj_view_form.tableView.setSortingEnabled(True)
        proj_view_form.tableView.resizeColumnsToContents()
        calc_konk()
        return_back1()


def return_back1():
    change_form.lineEdit_g1.clear()
    change_form.lineEdit_g7.clear()
    change_form.lineEdit_g5.clear()
    change_form.lineEdit_g6.clear()
    change_form.lineEdit_g8.clear()
    change_form.lineEdit_g9.clear()
    change_form.lineEdit_g10.clear()
    change_form.lineEdit_g11.clear()
    change_form.lineEdit_g1.setStyleSheet("color: black;")
    change_form.lineEdit_g7.setStyleSheet("color: black;")
    change_form.lineEdit_g5.setStyleSheet("color: black;")
    change_form.lineEdit_g6.setStyleSheet("color: black;")
    change_form.lineEdit_g8.setStyleSheet("color: black;")
    change_form.lineEdit_g9.setStyleSheet("color: black;")
    change_form.lineEdit_g10.setStyleSheet("color: black;")
    change_form.lineEdit_g11.setStyleSheet("color: black;")
    change_form.comboBox_cod_kon_change.clear()
    change_form.comboBox_cod_vuz_change.clear()
    change_form.comboBox_cod_nir_change.clear()
    change_window.close()
    proj_view_window.show()

def change():
    global res

    # заполняем комбобоксы
    g2 = set()
    g1=set()
    query = QSqlQuery("SELECT * FROM gr_proj")
    while query.next():
        g2.add(str(query.value(1))+'-'+str(query.value(0)))

    query = QSqlQuery("SELECT * FROM gr_proj")
    while query.next():
        g1.add(str(query.value(0)))

    g2=list(g2)
    #---------------------------------------------------------------------------------------------------------------------
    # запись в комбобок кодконкурса и код нира
    global comboBox_codkon
    comboBox_codkon=[]
    result = sorted([int(item) for item in g1])


    for j in range(1,16):
        for i in range(len(g2)):
            if int(g2[i][0:2])==j:
                comboBox_codkon.append(g2[i])


    result1 = [str(item) for item in g1]
    comboBox_codkon.sort()

    kod_con1 = ['01', '02', '03', '04', '05', '06', '07', '08', '09', '10', '11', '12', '13', '14', '15', '16', '17']
    def spisok(a, t):
        spis3 = []
        spis2 = [[], [], [], [], [], [], [], [], [], [], [], [], [], [], [], [], []]
        spis4 = []
        spis6 = []
        spis7 = []
        spis8 = []
        k = 0
        for i in kod_con1:
            for j in range(len(a)):
                if i == a[j][0:2]:
                    spis2[k].append((a[j]))
            spis3.append(spis2[k])
            k = k + 1
        for j in range(len(spis3[t])):
            spis4.append(int(spis3[t][j][3:]))
        spis5 = sorted(spis4)
        for i in range(len(spis5)):
            spis6.append(spis4.index(spis5[i]))
        for i in spis6:
            spis7.append(spis3[t][i])
        return spis7
    indexes = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16]
    global result3
    result3 = []

    def add_combo_box_change():
        add_combobox_codkon=[]
        print(change_form.comboBox_cod_kon_change.currentText()[0:2])
        print(result3)
        for i in range(len(result3)):
            if result3[i][0:2] == change_form.comboBox_cod_kon_change.currentText()[0:2]:
                add_combobox_codkon.append(result3[i])
        print(add_combobox_codkon)
        change_form.comboBox_cod_nir_change.clear()
        change_form.comboBox_cod_nir_change.addItems(add_combobox_codkon)


    for i in indexes:
        result3.append(spisok(comboBox_codkon, i))
    result3=list(chain.from_iterable(result3))

    change_form.comboBox_cod_kon_change.currentTextChanged.connect(add_combo_box_change)
    #------------------------------------------------------------------------------------------------------

    # заполняем код конкурса
    query = QSqlQuery("SELECT k2 , codkon FROM gr_konk")
    name_kon = set()
    while query.next():
        name_kon.add(str(query.value(1)) + '-' + str(query.value(0)))
    name_kon = list(name_kon)
    name_kon1=list(name_kon)

    change_form.comboBox_cod_kon_change.addItems(sorted(name_kon, key=lambda x: x[:3]))
    # заполняем код вуза и наименование вуза
    name_vuza = set()
    kod_vuza = set()
    name_vuza1=set()
    query = QSqlQuery("SELECT * FROM vuz")
    while query.next():
        name_vuza.add(str(query.value(3)) + '-' + str(query.value(0)))
        name_vuza1.add(str(query.value(3)))
    name_vuza = list(name_vuza)
    name_vuza1=sorted(name_vuza1, key=lambda x: x[:3])
    change_form.comboBox_cod_vuz_change.addItems(sorted(name_vuza, key=lambda x: x[:3]))
    b = []
    name_vuza=sorted(name_vuza, key=lambda x: x[:3])
    for i in range(len(name_vuza)):
        b.append((name_vuza[i].split("-")[0]))
    ###
    table_model = QSqlTableModel()
    global curr_index
    curr_index = proj_view_form.tableView.currentIndex()
    conn = sqlite3.connect(db_name)
    cur = conn.cursor()
    if curr_index.sibling(curr_index.row(), 1).data() != None:
        cur.execute(f'''SELECT * FROM gr_proj WHERE g1=={curr_index.sibling(curr_index.row(), 1).data()} and gr_proj.codkon=='{curr_index.sibling(curr_index.row(), 0).data()}' ''')
        res = list(cur.fetchone())
        indices = [i for i in range(0, len(comboBox_codkon)) if comboBox_codkon[i][0:2] == res[1]]
        kod_con=['01','02','03','04','05','06','07','08','09','10','11','12','13','14','15','16','17']
        change_form.comboBox_cod_kon_change.setCurrentIndex(kod_con.index(comboBox_codkon[indices[1]][0:2]))
        change_form.comboBox_cod_vuz_change.setCurrentIndex(b.index(res[3]))
        change_window.show()

        query = QSqlQuery(f"""SELECT codkon AS "Код конкурса",
                             g1 AS "Код НИР",
                             g8 AS "Руководитель НИР",
                             g6 AS "Наименование НИР",
                             z2 AS "Наименование вуза" ,
                             codvuz as "Код вуза",
                             g7 AS "Код по ГРНТИ",
                             g5 AS "Плановый объем",
                             g10 as"Ученое звание руководителя",
                             g11 as"Ученая степень руководителя"
                             FROM gr_proj WHERE g1=={curr_index.sibling(curr_index.row(), 1).data()} and gr_proj.codkon=='{curr_index.sibling(curr_index.row(), 0).data()}'""")

        table_model.setQuery(query)
        change_form.tableView_change.setModel(table_model)
        change_form.tableView_change.resizeColumnsToContents()
        ind = curr_index.sibling(curr_index.row(), 1).data()
        cod_kon_change=curr_index.sibling(curr_index.row(), 0).data()
        global spis

        spis = [result1, res, ind , cod_kon_change]
        change_form.lineEdit_g1.clear()
        change_form.lineEdit_g7.clear()
        change_form.lineEdit_g5.clear()
        change_form.lineEdit_g6.clear()
        change_form.lineEdit_g8.clear()
        change_form.lineEdit_g9.clear()
        change_form.lineEdit_g10.clear()
        change_form.lineEdit_g11.clear()
        change_form.lineEdit_g1.setText(str(spis[1][0]))
        change_form.lineEdit_g7.setText(str(spis[1][4]))
        change_form.lineEdit_g5.setText(str(spis[1][5]))
        change_form.lineEdit_g6.setText(str(spis[1][11]))
        change_form.lineEdit_g8.setText(str(spis[1][12]))
        change_form.lineEdit_g9.setText(str(spis[1][13]))
        change_form.lineEdit_g10.setText(str(spis[1][14]))
        change_form.lineEdit_g11.setText(str(spis[1][15]))
        calc_konk()


    else:
        print('ничего не выбрано')
        return_back1()


def open_proj_sort_window():
    # открывает окно показа таблицы конкурсов

    global query_sql_sort_save
    global query_sql_filter_save

    window_list[check_active_window()].close()
    proj_sort_window.showMaximized()

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    proj_sort_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                 proj.g1 AS "Код НИР",
                                 proj.g8 AS "Руководитель НИР",
                                 proj.g7 AS "Код по ГРНТИ",
                                 proj.z2 AS "Наименование вуза",
                                 proj.g5 AS "Плановый объем",
                                 proj.g6 AS "Наименование НИР",
                                 proj.g2 AS "Фактический объем гранта",
                                 proj.g21 as"Поквартальное финансирование",
                                 proj.g22 as"Поквартальное финансирование",
                                 proj.g23 as"Поквартальное финансирование",
                                 proj.g24 as"Поквартальное финансирование",
                                 proj.g9 as "Должность руководителя",
                                 proj.g10 as"Учёное звание руководителя",
                                 proj.g11 as"Учёная степень руководителя"
                          FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                          {cond_filter if query_sql_filter_save != {} else ""}
                          {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_sort_form.tableView.setSortingEnabled(True)
    proj_sort_form.tableView.setModel(table_model)
    proj_sort_form.tableView.resizeColumnsToContents()


def proj_sort_save():
    # сохраняет сортировку таблицы проектов

    global query_sql_sort_save
    global query_sql_filter_save
    global cond_sort
    global cond_filter


    combo_cond_sort_g1 = proj_sort_form.combo_box_order_g1.currentText()
    combo_cond_sort_codkon = proj_sort_form.combo_box_order_codkon.currentText()

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    proj_sort_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_sort_g1 = "ASC" if combo_cond_sort_g1 == "По возрастанию" else "DESC"
    cond_sort_codkon = "ASC" if combo_cond_sort_codkon == "По возрастанию" else "DESC"

    query_sql_sort_save = {'proj.g1': cond_sort_g1, 'proj.codkon': cond_sort_codkon}

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                 proj.g1 AS "Код НИР",
                                 proj.g8 AS "Руководитель НИР",
                                 proj.g7 AS "Код по ГРНТИ",
                                 proj.z2 AS "Наименование вуза",
                                 proj.g5 AS "Плановый объем",
                                 proj.g6 AS "Наименование НИР",
                                 proj.g2 AS "Фактический объем гранта",
                                 proj.g21 as"Поквартальное финансирование",
                                 proj.g22 as"Поквартальное финансирование",
                                 proj.g23 as"Поквартальное финансирование",
                                 proj.g24 as"Поквартальное финансирование",
                                 proj.g9 as "Должность руководителя",
                                 proj.g10 as"Учёное звание руководителя",
                                 proj.g11 as"Учёная степень руководителя"
                          FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                          {cond_filter if query_sql_filter_save != {} else ""}
                          {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_sort_form.tableView.setSortingEnabled(True)
    proj_sort_form.tableView.setModel(table_model)
    proj_sort_form.tableView.resizeColumnsToContents()


def proj_sort_reset():
    # удаляет сортировку таблицы проектов

    global query_sql_sort_save
    global query_sql_filter_save

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    proj_sort_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    query_sql_sort_save = {}

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                     proj.g1 AS "Код НИР",
                                     proj.g8 AS "Руководитель НИР",
                                     proj.g7 AS "Код по ГРНТИ",
                                     proj.z2 AS "Наименование вуза",
                                     proj.g5 AS "Плановый объем",
                                     proj.g6 AS "Наименование НИР",
                                     proj.g2 AS "Фактический объем гранта",
                                     proj.g21 as"Поквартальное финансирование",
                                     proj.g22 as"Поквартальное финансирование",
                                     proj.g23 as"Поквартальное финансирование",
                                     proj.g24 as"Поквартальное финансирование",
                                     proj.g9 as "Должность руководителя",
                                     proj.g10 as"Учёное звание руководителя",
                                     proj.g11 as"Учёная степень руководителя"
                          FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                          {cond_filter if query_sql_filter_save != {} else ""}
                          {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_sort_form.tableView.setSortingEnabled(True)
    proj_sort_form.tableView.setModel(table_model)
    proj_sort_form.tableView.resizeColumnsToContents()


def open_proj_filter_window():
    # открывает окно фильтрации таблицы проектов

    global query_sql_sort_save
    global query_sql_filter_save
    global query_sql_filter_temp

    window_list[check_active_window()].close()
    proj_filter_window.showMaximized()

    table_model = QSqlTableModel()

    if query_sql_filter_save == {}:
        query_sql_filter_temp = {}

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                         proj.g1 AS "Код НИР",
                                         proj.g8 AS "Руководитель НИР",
                                         proj.g7 AS "Код по ГРНТИ",
                                         proj.z2 AS "Наименование вуза",
                                         proj.g5 AS "Плановый объем",
                                         proj.g6 AS "Наименование НИР",
                                         proj.g2 AS "Фактический объем гранта",
                                         proj.g21 as"Поквартальное финансирование",
                                         proj.g22 as"Поквартальное финансирование",
                                         proj.g23 as"Поквартальное финансирование",
                                         proj.g24 as"Поквартальное финансирование",
                                         proj.g9 as "Должность руководителя",
                                         proj.g10 as"Учёное звание руководителя",
                                         proj.g11 as"Учёная степень руководителя"
                          FROM gr_proj AS proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                          {cond_filter if query_sql_filter_save != {} else ""}
                          {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    query = QSqlQuery(f"""SELECT konk.codkon, vuz.region AS region, vuz.oblname AS oblname, vuz.city AS city, vuz.z2 AS z2
                          FROM gr_konk AS konk
                          LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                          LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
    					  {cond_filter if query_sql_filter_save != {} else ""}""")

    set_codkon, set_region, set_oblname, set_city, set_z2 = set(), set(), set(), set(), set()

    while query.next():
        set_codkon.add(str(query.value(0)))
        set_region.add(str(query.value(1)))
        set_oblname.add(str(query.value(2)))
        set_city.add(str(query.value(3)))
        set_z2.add(str(query.value(4)))

    if '' in set_region and '' in set_oblname and '' in set_city and '' in set_z2:
        set_region.remove('')
        set_oblname.remove('')
        set_city.remove('')
        set_z2.remove('')

    proj_filter_form.combo_box_filter_codkon.clear()
    proj_filter_form.combo_box_filter_region.clear()
    proj_filter_form.combo_box_filter_oblname.clear()
    proj_filter_form.combo_box_filter_city.clear()
    proj_filter_form.combo_box_filter_z2.clear()

    proj_filter_form.combo_box_filter_codkon.addItems(["Нет"] + sorted(list(set_codkon)))
    proj_filter_form.combo_box_filter_region.addItems(["Нет"] + sorted(list(set_region)))
    proj_filter_form.combo_box_filter_oblname.addItems(["Нет"] + sorted(list(set_oblname)))
    proj_filter_form.combo_box_filter_city.addItems(["Нет"] + sorted(list(set_city)))
    proj_filter_form.combo_box_filter_z2.addItems(["Нет"] + sorted(list(set_z2)))

    if 'proj.codkon' in query_sql_filter_save:
        proj_filter_form.combo_box_filter_codkon.setCurrentIndex((["Нет"] + sorted(list(set_codkon))).index(query_sql_filter_save['proj.codkon']))

    if 'vuz.region' in query_sql_filter_save:
        proj_filter_form.combo_box_filter_region.setCurrentIndex((["Нет"] + sorted(list(set_region))).index(query_sql_filter_save['vuz.region']))

    if 'vuz.oblname' in query_sql_filter_save:
        proj_filter_form.combo_box_filter_oblname.setCurrentIndex((["Нет"] + sorted(list(set_oblname))).index(query_sql_filter_save['vuz.oblname']))

    if 'vuz.city' in query_sql_filter_save:
        proj_filter_form.combo_box_filter_city.setCurrentIndex((["Нет"] + sorted(list(set_city))).index(query_sql_filter_save['vuz.city']))

    if 'vuz.z2' in query_sql_filter_save:
        proj_filter_form.combo_box_filter_z2.setCurrentIndex((["Нет"] + sorted(list(set_z2))).index(query_sql_filter_save['vuz.z2']))


def proj_filter_codkon():
    # фильтрует по коду конкурса

    global query_sql_sort_save
    global query_sql_filter_temp

    table_model = QSqlTableModel()

    combo_box_filter_codkon = proj_filter_form.combo_box_filter_codkon.currentText()

    if combo_box_filter_codkon == "Нет" and 'proj.codkon' in query_sql_filter_temp:
        del query_sql_filter_temp['proj.codkon']
    elif combo_box_filter_codkon == "Нет":
        pass
    else:
        query_sql_filter_temp['proj.codkon'] = combo_box_filter_codkon

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                            proj.g1 AS "Код НИР",
                                            proj.g8 AS "Руководитель НИР",
                                            proj.g7 AS "Код по ГРНТИ",
                                            proj.z2 AS "Наименование вуза",
                                            proj.g5 AS "Плановый объем",
                                            proj.g6 AS "Наименование НИР",
                                            proj.g2 AS "Фактический объем гранта",
                                            proj.g21 as"Поквартальное финансирование",
                                            proj.g22 as"Поквартальное финансирование",
                                            proj.g23 as"Поквартальное финансирование",
                                            proj.g24 as"Поквартальное финансирование",
                                            proj.g9 as "Должность руководителя",
                                            proj.g10 as"Учёное звание руководителя",
                                            proj.g11 as"Учёная степень руководителя"
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz = vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}
                                     {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query_sql_filter)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    query_sql_filter = QSqlQuery(f"""SELECT vuz.region AS region, vuz.oblname AS oblname, vuz.city as city, vuz.z2 AS z2
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}""")

    set_region = set()
    set_oblname = set()
    set_city = set()
    set_z2 = set()

    query = QSqlQuery(query_sql_filter)

    while query.next():
        set_region.add(str(query.value(0)))
        set_oblname.add(str(query.value(1)))
        set_city.add(str(query.value(2)))
        set_z2.add(str(query.value(3)))

    proj_filter_form.combo_box_filter_region.clear()
    proj_filter_form.combo_box_filter_oblname.clear()
    proj_filter_form.combo_box_filter_city.clear()
    proj_filter_form.combo_box_filter_z2.clear()

    if len(set_region) == 1 and 'vuz.region' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_region.addItems(list(set_region))
    else:
        proj_filter_form.combo_box_filter_region.addItems(["Нет"] + sorted(list(set_region)))

    if len(set_oblname) == 1 and 'vuz.oblname' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_oblname.addItems(list(set_oblname))
    else:
        proj_filter_form.combo_box_filter_oblname.addItems(["Нет"] + sorted(list(set_oblname)))

    if len(set_city) == 1 and 'vuz.city' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_city.addItems(list(set_city))
    else:
        proj_filter_form.combo_box_filter_city.addItems(["Нет"] + sorted(list(set_city)))

    if len(set_z2) == 1 and 'vuz.z2' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_z2.addItems(list(set_z2))
    else:
        proj_filter_form.combo_box_filter_z2.addItems(["Нет"] + sorted(list(set_z2)))


def proj_filter_region():
    # фильтрует по региону конкурса

    global query_sql_sort_save
    global query_sql_filter_temp

    table_model = QSqlTableModel()

    combo_box_filter_region = proj_filter_form.combo_box_filter_region.currentText()

    if combo_box_filter_region == "Нет" and 'vuz.region' in query_sql_filter_temp:
        del query_sql_filter_temp['vuz.region']
    elif combo_box_filter_region == "Нет":
        pass
    else:
        query_sql_filter_temp['vuz.region'] = combo_box_filter_region

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                            proj.g1 AS "Код НИР",
                                            proj.g8 AS "Руководитель НИР",
                                            proj.g7 AS "Код по ГРНТИ",
                                            proj.z2 AS "Наименование вуза",
                                            proj.g5 AS "Плановый объем",
                                            proj.g6 AS "Наименование НИР",
                                            proj.g2 AS "Фактический объем гранта",
                                            proj.g21 as"Поквартальное финансирование",
                                            proj.g22 as"Поквартальное финансирование",
                                            proj.g23 as"Поквартальное финансирование",
                                            proj.g24 as"Поквартальное финансирование",
                                            proj.g9 as "Должность руководителя",
                                            proj.g10 as"Учёное звание руководителя",
                                            proj.g11 as"Учёная степень руководителя"
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz = vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}
                                     {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query_sql_filter)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS codkon, vuz.oblname AS oblname, vuz.city as city, vuz.z2 AS z2
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}""")

    set_codkon = set()
    set_oblname = set()
    set_city = set()
    set_z2 = set()

    query = QSqlQuery(query_sql_filter)

    while query.next():
        set_codkon.add(str(query.value(0)))
        set_oblname.add(str(query.value(1)))
        set_city.add(str(query.value(2)))
        set_z2.add(str(query.value(3)))

    proj_filter_form.combo_box_filter_codkon.clear()
    proj_filter_form.combo_box_filter_oblname.clear()
    proj_filter_form.combo_box_filter_city.clear()
    proj_filter_form.combo_box_filter_z2.clear()

    if len(set_codkon) == 1 and 'proj.codkon' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_codkon.addItems(list(set_codkon))
    else:
        proj_filter_form.combo_box_filter_codkon.addItems(["Нет"] + sorted(list(set_codkon)))

    if len(set_oblname) == 1 and 'vuz.oblname' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_oblname.addItems(list(set_oblname))
    else:
        proj_filter_form.combo_box_filter_oblname.addItems(["Нет"] + sorted(list(set_oblname)))

    if len(set_city) == 1 and 'vuz.city' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_city.addItems(list(set_city))
    else:
        proj_filter_form.combo_box_filter_city.addItems(["Нет"] + sorted(list(set_city)))

    if len(set_z2) == 1 and 'vuz.z2' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_z2.addItems(list(set_z2))
    else:
        proj_filter_form.combo_box_filter_z2.addItems(["Нет"] + sorted(list(set_z2)))


def proj_filter_oblname():
    # фильтрует по субъекту РФ конкурса

    global query_sql_sort_save
    global query_sql_filter_temp

    table_model = QSqlTableModel()

    combo_box_filter_oblname = proj_filter_form.combo_box_filter_oblname.currentText()

    if combo_box_filter_oblname == "Нет" and 'vuz.oblname' in query_sql_filter_temp:
        del query_sql_filter_temp['vuz.oblname']
    elif combo_box_filter_oblname == "Нет":
        pass
    else:
        query_sql_filter_temp['vuz.oblname'] = combo_box_filter_oblname

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                            proj.g1 AS "Код НИР",
                                            proj.g8 AS "Руководитель НИР",
                                            proj.g7 AS "Код по ГРНТИ",
                                            proj.z2 AS "Наименование вуза",
                                            proj.g5 AS "Плановый объем",
                                            proj.g6 AS "Наименование НИР",
                                            proj.g2 AS "Фактический объем гранта",
                                            proj.g21 as"Поквартальное финансирование",
                                            proj.g22 as"Поквартальное финансирование",
                                            proj.g23 as"Поквартальное финансирование",
                                            proj.g24 as"Поквартальное финансирование",
                                            proj.g9 as "Должность руководителя",
                                            proj.g10 as"Учёное звание руководителя",
                                            proj.g11 as"Учёная степень руководителя"
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz = vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}
                                     {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query_sql_filter)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS codkon, vuz.region AS region, vuz.city as city, vuz.z2 AS z2
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}""")

    set_codkon = set()
    set_region = set()
    set_city = set()
    set_z2 = set()

    query = QSqlQuery(query_sql_filter)

    while query.next():
        set_codkon.add(str(query.value(0)))
        set_region.add(str(query.value(1)))
        set_city.add(str(query.value(2)))
        set_z2.add(str(query.value(3)))

    proj_filter_form.combo_box_filter_codkon.clear()
    proj_filter_form.combo_box_filter_region.clear()
    proj_filter_form.combo_box_filter_city.clear()
    proj_filter_form.combo_box_filter_z2.clear()

    if len(set_codkon) == 1 and 'proj.codkon' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_codkon.addItems(list(set_codkon))
    else:
        proj_filter_form.combo_box_filter_codkon.addItems(["Нет"] + sorted(list(set_codkon)))

    if len(set_region) == 1 and 'vuz.region' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_region.addItems(list(set_region))
    else:
        proj_filter_form.combo_box_filter_region.addItems(["Нет"] + sorted(list(set_region)))

    if len(set_city) == 1 and 'vuz.city' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_city.addItems(list(set_city))
    else:
        proj_filter_form.combo_box_filter_city.addItems(["Нет"] + sorted(list(set_city)))

    if len(set_z2) == 1 and 'vuz.z2' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_z2.addItems(list(set_z2))
    else:
        proj_filter_form.combo_box_filter_z2.addItems(["Нет"] + sorted(list(set_z2)))


def proj_filter_city():
    # фильтрует по городу конкурса

    global query_sql_sort_save
    global query_sql_filter_temp

    table_model = QSqlTableModel()

    combo_box_filter_city = proj_filter_form.combo_box_filter_city.currentText()

    if combo_box_filter_city == "Нет" and 'vuz.city' in query_sql_filter_temp:
        del query_sql_filter_temp['vuz.city']
    elif combo_box_filter_city == "Нет":
        pass
    else:
        query_sql_filter_temp['vuz.city'] = combo_box_filter_city

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                            proj.g1 AS "Код НИР",
                                            proj.g8 AS "Руководитель НИР",
                                            proj.g7 AS "Код по ГРНТИ",
                                            proj.z2 AS "Наименование вуза",
                                            proj.g5 AS "Плановый объем",
                                            proj.g6 AS "Наименование НИР",
                                            proj.g2 AS "Фактический объем гранта",
                                            proj.g21 as"Поквартальное финансирование",
                                            proj.g22 as"Поквартальное финансирование",
                                            proj.g23 as"Поквартальное финансирование",
                                            proj.g24 as"Поквартальное финансирование",
                                            proj.g9 as "Должность руководителя",
                                            proj.g10 as"Учёное звание руководителя",
                                            proj.g11 as"Учёная степень руководителя"
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz = vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}
                                     {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query_sql_filter)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS codkon, vuz.region AS region, vuz.oblname AS oblname, vuz.z2 AS z2
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}""")

    set_codkon = set()
    set_region = set()
    set_oblname = set()
    set_z2 = set()

    query = QSqlQuery(query_sql_filter)

    while query.next():
        set_codkon.add(str(query.value(0)))
        set_region.add(str(query.value(1)))
        set_oblname.add(str(query.value(2)))
        set_z2.add(str(query.value(3)))

    proj_filter_form.combo_box_filter_codkon.clear()
    proj_filter_form.combo_box_filter_region.clear()
    proj_filter_form.combo_box_filter_oblname.clear()
    proj_filter_form.combo_box_filter_z2.clear()

    if len(set_codkon) == 1 and 'proj.codkon' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_codkon.addItems(list(set_codkon))
    else:
        proj_filter_form.combo_box_filter_codkon.addItems(["Нет"] + sorted(list(set_codkon)))

    if len(set_region) == 1 and 'vuz.region' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_region.addItems(list(set_region))
    else:
        proj_filter_form.combo_box_filter_region.addItems(["Нет"] + sorted(list(set_region)))

    if len(set_oblname) == 1 and 'vuz.oblname' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_oblname.addItems(list(set_oblname))
    else:
        proj_filter_form.combo_box_filter_oblname.addItems(["Нет"] + sorted(list(set_oblname)))

    if len(set_z2) == 1 and 'vuz.z2' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_z2.addItems(list(set_z2))
    else:
        proj_filter_form.combo_box_filter_z2.addItems(["Нет"] + sorted(list(set_z2)))


def proj_filter_z2():
    # фильтрует по вузу конкурса

    global query_sql_sort_save
    global query_sql_filter_temp

    table_model = QSqlTableModel()

    combo_box_filter_z2 = proj_filter_form.combo_box_filter_z2.currentText()

    if combo_box_filter_z2 == "Нет" and 'vuz.z2' in query_sql_filter_temp:
        del query_sql_filter_temp['vuz.z2']
    elif combo_box_filter_z2 == "Нет":
        pass
    else:
        query_sql_filter_temp['vuz.z2'] = combo_box_filter_z2

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                            proj.g1 AS "Код НИР",
                                            proj.g8 AS "Руководитель НИР",
                                            proj.g7 AS "Код по ГРНТИ",
                                            proj.z2 AS "Наименование вуза",
                                            proj.g5 AS "Плановый объем",
                                            proj.g6 AS "Наименование НИР",
                                            proj.g2 AS "Фактический объем гранта",
                                            proj.g21 as"Поквартальное финансирование",
                                            proj.g22 as"Поквартальное финансирование",
                                            proj.g23 as"Поквартальное финансирование",
                                            proj.g24 as"Поквартальное финансирование",
                                            proj.g9 as "Должность руководителя",
                                            proj.g10 as"Учёное звание руководителя",
                                            proj.g11 as"Учёная степень руководителя"
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz = vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}
                                     {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query_sql_filter)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    query_sql_filter = QSqlQuery(f"""SELECT proj.codkon AS codkon, vuz.region AS region, vuz.oblname AS oblname, vuz.city AS city
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz == vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}""")

    set_codkon = set()
    set_region = set()
    set_oblname = set()
    set_city = set()

    query = QSqlQuery(query_sql_filter)

    while query.next():
        set_codkon.add(str(query.value(0)))
        set_region.add(str(query.value(1)))
        set_oblname.add(str(query.value(2)))
        set_city.add(str(query.value(3)))

    proj_filter_form.combo_box_filter_codkon.clear()
    proj_filter_form.combo_box_filter_region.clear()
    proj_filter_form.combo_box_filter_oblname.clear()
    proj_filter_form.combo_box_filter_city.clear()

    if len(set_codkon) == 1 and 'proj.codkon' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_codkon.addItems(list(set_codkon))
    else:
        proj_filter_form.combo_box_filter_codkon.addItems(["Нет"] + sorted(list(set_codkon)))

    if len(set_region) == 1 and 'vuz.region' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_region.addItems(list(set_region))
    else:
        proj_filter_form.combo_box_filter_region.addItems(["Нет"] + sorted(list(set_region)))

    if len(set_oblname) == 1 and 'vuz.oblname' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_oblname.addItems(list(set_oblname))
    else:
        proj_filter_form.combo_box_filter_oblname.addItems(["Нет"] + sorted(list(set_oblname)))

    if len(set_city) == 1 and 'vuz.city' in query_sql_filter_temp:
        proj_filter_form.combo_box_filter_city.addItems(list(set_city))
    else:
        proj_filter_form.combo_box_filter_city.addItems(["Нет"] + sorted(list(set_city)))


def proj_filter_save():
    # сохранение фильтра

    global query_sql_filter_save
    global query_sql_filter_temp

    query_sql_filter_save = copy.deepcopy(query_sql_filter_temp)

    calc_konk()


def proj_filter_reset():
    # удаление фильтра

    global query_sql_sort_save
    global query_sql_filter_save
    global query_sql_filter_temp

    query_sql_filter_temp = {}
    query_sql_filter_save = {}

    table_model = QSqlTableModel()

    cond_sort = 'ORDER BY' + ' ' + ', '.join(f"{field} {value}" for field, value in zip(query_sql_sort_save.keys(), query_sql_sort_save.values()))
    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT proj.codkon AS "Код конкурса",
                                            proj.g1 AS "Код НИР",
                                            proj.g8 AS "Руководитель НИР",
                                            proj.g7 AS "Код по ГРНТИ",
                                            proj.z2 AS "Наименование вуза",
                                            proj.g5 AS "Плановый объем",
                                            proj.g6 AS "Наименование НИР",
                                            proj.g2 AS "Фактический объем гранта",
                                            proj.g21 as"Поквартальное финансирование",
                                            proj.g22 as"Поквартальное финансирование",
                                            proj.g23 as"Поквартальное финансирование",
                                            proj.g24 as"Поквартальное финансирование",
                                            proj.g9 as "Должность руководителя",
                                            proj.g10 as"Учёное звание руководителя",
                                            proj.g11 as"Учёная степень руководителя"
                                     FROM gr_proj as proj INNER JOIN vuz ON proj.codvuz = vuz.codvuz
                                     {cond_filter if query_sql_filter_temp != {} else ""}
                                     {cond_sort if query_sql_sort_save != {} else ""}""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    proj_filter_form.tableView.setSortingEnabled(True)
    proj_filter_form.tableView.setModel(table_model)
    proj_filter_form.tableView.resizeColumnsToContents()

    calc_konk()

    query = QSqlQuery(f"""SELECT konk.codkon, vuz.region AS region, vuz.oblname AS oblname, vuz.city AS city, vuz.z2 AS z2
                          FROM gr_konk AS konk
                          LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                          LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
    					  {cond_filter if query_sql_filter_save != {} else ""}""")

    set_codkon, set_region, set_oblname, set_city, set_z2 = set(), set(), set(), set(), set()

    while query.next():
        set_codkon.add(str(query.value(0)))
        set_region.add(str(query.value(1)))
        set_oblname.add(str(query.value(2)))
        set_city.add(str(query.value(3)))
        set_z2.add(str(query.value(4)))

    if '' in set_region and '' in set_oblname and '' in set_city and '' in set_z2:
        set_region.remove('')
        set_oblname.remove('')
        set_city.remove('')
        set_z2.remove('')

    proj_filter_form.combo_box_filter_codkon.clear()
    proj_filter_form.combo_box_filter_region.clear()
    proj_filter_form.combo_box_filter_oblname.clear()
    proj_filter_form.combo_box_filter_city.clear()
    proj_filter_form.combo_box_filter_z2.clear()

    proj_filter_form.combo_box_filter_codkon.addItems(["Нет"] + sorted(list(set_codkon)))
    proj_filter_form.combo_box_filter_region.addItems(["Нет"] + sorted(list(set_region)))
    proj_filter_form.combo_box_filter_oblname.addItems(["Нет"] + sorted(list(set_oblname)))
    proj_filter_form.combo_box_filter_city.addItems(["Нет"] + sorted(list(set_city)))
    proj_filter_form.combo_box_filter_z2.addItems(["Нет"] + sorted(list(set_z2)))


def proj_filter_check():
    print('temp',query_sql_filter_temp)
    print('save',query_sql_filter_save)


def open_konk_window():
    # открывает окно показа таблицы конкурсов

    global query_sql_filter_save

    if check_active_window() is not None:
        window_list[check_active_window()].close()
    konk_view_window.showMaximized()

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    konk_view_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    calc_konk()

    query = QSqlQuery(f"""SELECT k2 AS "Название конкурса",
                                 codkon AS "Код конкурса",
                                 k12 AS "Плановый объём",
    			                 k4 AS "Фактический объём",
    			                 k41 AS "Финанисирование в 1 квартале",
    			                 k42 AS "Финанисирование во 2 квартале",
    			                 k43 AS "Финанисирование в 3 квартале",
    			                 k44 AS "Финанисирование в 4 квартале",
                                 npr AS "Количество НИР"
                          FROM gr_konk""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    konk_view_form.tableView.setSortingEnabled(True)
    konk_view_form.tableView.setModel(table_model)
    konk_view_form.tableView.resizeColumnsToContents()


def calc_konk():
    # расчёт таблицы конкурсов

    global query_sql_filter_save

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT konk.k2 AS k2, 
                                     konk.codkon AS codkon,
                                     ifnull(SUM(proj.g5), 0) AS k12,
        			                 ifnull(SUM(proj.g2), 0) AS k4, 
        			                 ifnull(SUM(proj.g21), 0) AS k41, 
        			                 ifnull(SUM(proj.g22), 0) AS k42, 
        			                 ifnull(SUM(proj.g23), 0) AS k43, 
        			                 ifnull(SUM(proj.g24), 0) AS k44, 
                                     COUNT(proj.codkon) AS npr
                              FROM gr_konk AS konk
                              LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                              LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
    						  {cond_filter if query_sql_filter_save != {} else ""}
    						  GROUP BY konk.codkon""")

    list_k2, list_codkon, list_k12, list_k4, list_k41, list_k42, list_k43, list_k44, list_npr = list(), list(), list(), list(), list(), list(), list(), list(), list()

    while query.next():
        list_k2.append(str(query.value(0)))
        list_codkon.append(str(query.value(1)))
        list_k12.append(query.value(2))
        list_k4.append(query.value(3))
        list_k41.append(query.value(4))
        list_k42.append(query.value(5))
        list_k43.append(query.value(6))
        list_k44.append(query.value(7))
        list_npr.append(query.value(8))

    record_list = list(zip(list_k12, list_k4, list_k41, list_k42, list_k43, list_k44, list_npr, list_k2, list_codkon))

    conn = sqlite3.connect(db_name)
    cur = conn.cursor()

    sqlite_update_query = """UPDATE gr_konk SET
                             k12 = ?,
                             k4 = ?,
                             k41 = ?,
                             k42 = ?,
                             k43 = ?,
                             k44 = ?,
                             npr = ?
                             WHERE k2 = ? AND codkon = ?"""

    cur.executemany(sqlite_update_query, record_list)

    conn.commit()
    cur.close()
    conn.close()


def open_vuz_window():
    # открывает окно показа таблицы вузов

    global query_sql_filter_save

    window_list[check_active_window()].close()
    vuz_view_window.showMaximized()

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    vuz_view_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT DISTINCT vuz.codvuz AS "Код вуза",
                                 vuz.z1 AS "Наименование вуза",
                                 vuz.z1full AS "Полное юридическое наимнование вуза",
                                 vuz.z2 AS "Сокращённое наименование вуза",
                                 vuz.region AS "Федеральный округ",
                                 vuz.city AS "Город",
                                 vuz.status AS "Статус",
                                 vuz.obl AS "Код субъекта РФ",
                                 vuz.oblname AS "Субъект РФ",
                                 vuz.gr_ved AS "Группа ведущих вузов",
                                 vuz.prof AS "Профиль вуза"
						  FROM vuz LEFT JOIN gr_proj AS proj ON vuz.codvuz = proj.codvuz
						  {cond_filter if query_sql_filter_save != {} else ""}""")

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    vuz_view_form.tableView.setSortingEnabled(True)
    vuz_view_form.tableView.setModel(table_model)
    vuz_view_form.tableView.resizeColumnsToContents()

# сохраняем в таблицу по анализу вузов

def add_doc_analys(word_add_analys,filter_analys,name_doc):
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading(f'{name_doc}',0)

    # Table data in a form of list

    # Creating a table object
    table = doc.add_table(rows=1, cols=4)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Наименование ВУЗа'
    row[1].text = 'Количество НИР'
    row[2].text ='Плановый объем'
    row[3].text ='Количество конкурсов ,в которых участвует ВУЗ'

    # Adding data from the list to the table
    for name, kol_nir , plan_v ,kol_con in word_add_analys:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = name
        row[1].text = kol_nir
        row[2].text = plan_v
        row[3].text= kol_con
    # добавление в конце надписи фильтр

    date_object=datetime.today()
    doc.add_paragraph('Создан: '+str(date_object)[:-7])

    doc.add_paragraph(str(filter_analys))

    # Now save the document to a location
    doc.save(f'{name_doc}.docx')



def open_analysis_vuz_window():
    # открывает окно анализа по вузам
    global query_sql_filter_save

    window_list[check_active_window()].close()
    analysis_window.showMaximized()
    analysis_window.setWindowTitle('Сопровождение конкурсов на соискание грантов — анализ по вузам')

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    analysis_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    query = QSqlQuery(f"""WITH temp_cte AS(
                            SELECT vuz.z2 AS z2_cte,
                            COUNT(proj.g1) AS g1_cte,
                            ifnull(SUM(proj.g5), 0) AS g5_cte,
                            COUNT(DISTINCT proj.codkon) AS codkon_cte
                            FROM vuz LEFT JOIN gr_proj AS proj ON vuz.codvuz = proj.codvuz
                            {cond_filter if query_sql_filter_save != {} else ""}
                            GROUP BY proj.codvuz)
                         SELECT z2_cte AS "Наименование вуза",
                                g1_cte AS "Количество НИР",
                                g5_cte AS "Плановый объём",
                                codkon_cte AS "Количество конкурсов в которых учавствует вуз"
                         FROM temp_cte
                         WHERE g1_cte > 0
                         UNION ALL
                         SELECT NULL, NULL, NULL, NULL
                         UNION ALL
                         SELECT 'Итого',
                         		ifnull(SUM(g1_cte), 0),
                         		ifnull(SUM(g5_cte), 0),
                         		NULL
                         FROM temp_cte""")

    name_doc = 'Анализ по ВУЗам'
    filter_analys = filter_label
    if filter_analys == '':
        filter_analys = 'Условия фильтрации: Отсутсвуют'
    else:
        filter_analys = 'Условия фильтрации:' + str(filter_label)
    word_add_analys=[]
    spis = []
    while query.next():
        spis.append(str(query.value(0)))
        spis.append(str(query.value(1)))
        spis.append(str(query.value(2)))
        spis.append(str(query.value(3)))
        word_add_analys.append(spis)
        spis = []

    analysis_form.push_button_print.clicked.connect(partial(add_doc_analys,word_add_analys,filter_analys,name_doc))

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    analysis_form.tableView.setSortingEnabled(True)
    analysis_form.tableView.setModel(table_model)
    analysis_form.tableView.resizeColumnsToContents()


def open_analysis_konk_window():
    # открывает окно анализа по конкурсам

    global query_sql_filter_save

    window_list[check_active_window()].close()
    analysis_window.showMaximized()
    analysis_window.setWindowTitle('Сопровождение конкурсов на соискание грантов — анализ по конкурсам')
    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in
                             zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    analysis_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""WITH temp_cte AS(
                            SELECT konk.k2 AS k2_cte,
                            COUNT(proj.g1) AS g1_cte,
                            ifnull(SUM(proj.g5), 0) AS g5_cte,
                            COUNT(DISTINCT proj.codvuz) AS codvuz_cte
                            FROM gr_konk AS konk
                            LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                            LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
                            {cond_filter if query_sql_filter_save != {} else ""}
                            GROUP BY konk.codkon)
                         SELECT k2_cte AS "Название конкурса",
                                g1_cte AS "Количество НИР",
                                g5_cte AS "Плановый объём",
                                codvuz_cte AS "Количество вузов которые учавствуют в конкурсе"
                         FROM temp_cte
                         UNION ALL
                         SELECT NULL, NULL, NULL, NULL
                         UNION ALL
                         SELECT 'Итого',
                         		ifnull(SUM(g1_cte), 0),
                         		NULL,
                         		ifnull(SUM(codvuz_cte), 0)
                         FROM temp_cte""")

    name_doc = 'Анализ по конкурсам'
    filter_analys = filter_label
    if filter_analys == '':
        filter_analys = 'Условия фильтрации: Отсутствуют'
    else:
        filter_analys = 'Условия фильтрации:' + str(filter_label)
    word_add_analys = []
    spis = []
    while query.next():
        spis.append(str(query.value(0)))
        spis.append(str(query.value(1)))
        spis.append(str(query.value(2)))
        spis.append(str(query.value(3)))
        word_add_analys.append(spis)
        spis = []

    analysis_form.push_button_print.clicked.connect(partial(add_doc_analys, word_add_analys, filter_analys, name_doc))

    table_model.setQuery(query)
    analysis_form.tableView.setSortingEnabled(True)
    analysis_form.tableView.setModel(table_model)
    analysis_form.tableView.resizeColumnsToContents()

def open_analysis_subj_window():
    # открывает окно анализа по субъектам РФ

    global query_sql_filter_save

    window_list[check_active_window()].close()
    analysis_window.showMaximized()
    analysis_window.setWindowTitle('Сопровождение конкурсов на соискание грантов — анализ по субъектам РФ')

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    analysis_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")

    table_model = QSqlTableModel()

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""WITH temp_cte AS(
                            SELECT vuz.oblname AS oblname_cte,
                            COUNT(proj.g1) AS g1_cte,
                            ifnull(SUM(proj.g5), 0) AS g5_cte,
							COUNT(DISTINCT vuz.codvuz) AS codvuz_cte,
                            COUNT(DISTINCT proj.codkon) AS codkon_cte
                            FROM vuz LEFT JOIN gr_proj AS proj ON vuz.codvuz = proj.codvuz
                            {cond_filter if query_sql_filter_save != {} else ""}
                            GROUP BY vuz.oblname)
                         SELECT oblname_cte AS "Субъект федерации",
                                g1_cte AS "Количество НИР",
                                g5_cte AS "Плановый объём",
								codvuz_cte AS "Количество вузов в субъекте",
                                codkon_cte AS "Количество конкурсов в субъекте"
                         FROM temp_cte
                         WHERE g1_cte > 0
                         UNION ALL
                         SELECT NULL, NULL, NULL, NULL, NULL
                         UNION ALL
                         SELECT 'Итого',
                         		ifnull(SUM(g1_cte), 0),
                         		ifnull(SUM(g5_cte), 0),
								ifnull(SUM(codvuz_cte), 0),
                         		NULL
                         FROM temp_cte""")

    name_doc = 'Анализ по субьектам'
    filter_analys = filter_label
    if filter_analys == '':
        filter_analys = 'Условия фильтрации: Отсутсвует'
    else:
        filter_analys = 'Условия фильтрации:' + str(filter_label)
    word_add_analys = []
    spis = []
    while query.next():
        spis.append(str(query.value(0)))
        spis.append(str(query.value(1)))
        spis.append(str(query.value(2)))
        spis.append(str(query.value(3)))
        word_add_analys.append(spis)
        spis = []

    analysis_form.push_button_print.clicked.connect(partial(add_doc_analys, word_add_analys, filter_analys, name_doc))

    table_model.setQuery(query)

    while table_model.canFetchMore():
        table_model.fetchMore()

    analysis_form.tableView.setSortingEnabled(True)
    analysis_form.tableView.setModel(table_model)
    analysis_form.tableView.resizeColumnsToContents()

def open_fin_window():
    # открывает окно финансирования

    # надо сделать ограничения на проценты после имеющегося финансирования
    # надо сделать ограничения на сумму после имеющегося финансирования

    def calc_fin_sum():
        # расчёт суммы по процентам
        if fin_form.percentfin.text()!='':
            fin_percent = round(ast.literal_eval(fin_form.percentfin.text()) / 100,3) if isinstance(ast.literal_eval(fin_form.percentfin.text()), float) \
                else ast.literal_eval(fin_form.percentfin.text()) / 100
            fin_form.percentfin.setText(str(round(fin_percent*100,1)))
        else:
            fin_percent = 0
        if fin_percent > (sum_plan-sum_fact)/sum_plan:
            fin_percent = round((sum_plan-sum_fact)/sum_plan,3)
            fin_form.percentfin.setText(str(fin_percent * 100))
        fin_sum = fin_percent * sum_plan
        fin_form.sumfin.setText(str(fin_sum))

    def calc_fin_percent():
        # расчёт процентов по сумме
        fin_sum = ast.literal_eval(fin_form.sumfin.text()) if fin_form.sumfin.text()!='' else 0
        if fin_sum > sum_plan - sum_fact:
            fin_sum = sum_plan - sum_fact
            fin_form.sumfin.setText(str(fin_sum))
        fin_percent = round(fin_sum / sum_plan * 100,1)
        fin_form.percentfin.setText(str(fin_percent))

    global query_sql_filter_save

    window_list[check_active_window()].close()
    fin_window.showMaximized()

    try:
        fin_form.tableView.model().clear()
    except:
        pass

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))

    query = QSqlQuery(f"""SELECT ifnull(SUM(proj.g5), 0) AS sum_plan, ifnull(SUM(proj.g2), 0) AS sum_fact
                          FROM gr_konk AS konk
                          LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                          LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
    					  {cond_filter if query_sql_filter_save != {} else ""}""")

    while query.next():
        sum_plan = query.value(0)
        sum_fact = query.value(1)

    filter_label = ', '.join(f"{translate_dict_filter[field]}: {value}" for field, value in zip(query_sql_filter_temp.keys(), query_sql_filter_temp.values()))
    fin_form.filterlabel.setText(filter_label if filter_label != '' else "Отсутствуют")
    fin_form.planlabel.setText(str(sum_plan) + ' ' + 'руб.')
    fin_form.factlabel.setText(str(sum_fact) + ' ' + 'руб.')

    fin_form.sumfin.editingFinished.connect(calc_fin_percent)
    fin_form.percentfin.editingFinished.connect(calc_fin_sum)

    fin_form.push_button_save_fin.setEnabled(False)
    fin_form.push_button_cancel_fin.setEnabled(False)
    fin_form.push_button_print_fin.setEnabled(False)




def doc_add_fin(word_add_analys,number_kvartala,name_doc, filter):
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading(f'{name_doc}', 0)

    # Table data in a form of list

    # Creating a table object
    table = doc.add_table(rows=1, cols=2)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Наименование ВУЗа'
    row[1].text = f'Финансирование за {number_kvartala}'

    # Adding data from the list to the table
    for name, kol_nir  in word_add_analys:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = name
        row[1].text = kol_nir

    date_object = datetime.today()
    doc.add_paragraph('Создан: ' + str(date_object)[:-7])

    doc.add_paragraph(str(filter))
    # добавление в конце надписи фильтр


    # Now save the document to a location
    doc.save(f'{name_doc}.docx')

def calc_fin():
    # расчёт финансирования

    global query_sql_filter_save

    table_model = QSqlTableModel()

    cond_filter = 'WHERE' + ' ' + ' AND '.join(f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_save.keys(), query_sql_filter_save.values()))

    if fin_form.percentfin.text()!='' and fin_form.sumfin.text() != '' and fin_form.quartcombo.currentText() != '':

        quarter_dict = {'I квартал': '1', 'II квартал': '2', 'III квартал': '3', 'IV квартал': '4'}
        quarter = quarter_dict[fin_form.quartcombo.currentText()]

        fin_percent = str(round(ast.literal_eval(fin_form.percentfin.text()) / 100, 2) if isinstance(ast.literal_eval(fin_form.percentfin.text()), float) \
            else ast.literal_eval(fin_form.percentfin.text()) / 100)

        query = QSqlQuery(f"""WITH temp_cte AS(
                                SELECT vuz.z2 AS z2_cte,
                                ifnull(SUM(proj.g5) * {fin_percent}, 0) AS g2{quarter}_cte
                                FROM vuz LEFT JOIN gr_proj AS proj ON vuz.codvuz = proj.codvuz
                                {cond_filter if query_sql_filter_save != {} else ""}
                                GROUP BY proj.codvuz)
                             SELECT z2_cte AS "Наименование вуза",
                                    g2{quarter}_cte AS "Финансирование за {fin_form.quartcombo.currentText()}"
                             FROM temp_cte
                             WHERE g2{quarter}_cte > 0
                             UNION ALL
                             SELECT NULL, NULL
                             UNION ALL
                             SELECT 'Итого',
                                    ifnull(SUM(g2{quarter}_cte), 0)
                             FROM temp_cte""")
        filter='Условия фильтрации: '+str(fin_form.filterlabel.text())

        name_doc = 'Ведомость по финансированию ВУЗов'
        number_kvartala=fin_form.quartcombo.currentText()
        word_add_analys = []
        spis = []
        while query.next():
            spis.append(str(query.value(0)))
            spis.append(str(query.value(1)))
            word_add_analys.append(spis)
            spis = []
        print(word_add_analys)
        fin_form.push_button_print_fin.clicked.connect(partial(doc_add_fin,word_add_analys,number_kvartala,name_doc,filter))
        table_model.setQuery(query)

        while table_model.canFetchMore():
            table_model.fetchMore()

        fin_form.tableView.setSortingEnabled(True)
        fin_form.tableView.setModel(table_model)
        fin_form.tableView.resizeColumnsToContents()

        fin_form.push_button_save_fin.setEnabled(True)
        fin_form.push_button_cancel_fin.setEnabled(True)

    else:
        pass


def save_fin():
    # сохранение финансирования

    conn = sqlite3.connect(db_name)
    cur = conn.cursor()

    sqlite_drop_temp_table_query = """DROP TABLE IF EXISTS temp_proj"""
    cur.execute(sqlite_drop_temp_table_query)
    conn.commit()

    sqlite_create_temp_table_query = """CREATE TABLE temp_proj AS SELECT * FROM gr_proj"""
    cur.execute(sqlite_create_temp_table_query)
    conn.commit()

    cur.close()
    conn.close()

    global query_sql_filter_save

    cond_filter = 'WHERE' + ' ' + ' AND '.join(
        f"{field} == \"{value}\"" for field, value in zip(query_sql_filter_save.keys(), query_sql_filter_save.values()))

    quarter_dict = {'I квартал': '1', 'II квартал': '2', 'III квартал': '3', 'IV квартал': '4'}
    quarter = quarter_dict[fin_form.quartcombo.currentText()]

    fin_percent = str(round(ast.literal_eval(fin_form.percentfin.text()) / 100, 2) if isinstance(
        ast.literal_eval(fin_form.percentfin.text()), float) \
                          else ast.literal_eval(fin_form.percentfin.text()) / 100)

    query = QSqlQuery(f"""SELECT proj.g1 AS g1_cte,
                                 proj.codkon AS codkon_cte,
                                 proj.g5 * {fin_percent} AS g2{quarter}_cte
                          FROM vuz LEFT JOIN gr_proj AS proj ON vuz.codvuz = proj.codvuz
                          {cond_filter if query_sql_filter_save != {} else ""}""")

    list_g1, list_codkon, list_g_quarter = list(), list(), list()

    while query.next():
        list_g1.append(query.value(0))
        list_codkon.append(query.value(1))
        list_g_quarter.append(query.value(2))

    record_list_quarter = list(zip(list_g_quarter, list_g1, list_codkon))
    record_list_g2 = list(zip(list_g1, list_codkon))

    conn = sqlite3.connect(db_name)
    cur = conn.cursor()

    sqlite_update_query = f"""UPDATE gr_proj SET
                             g2{quarter} = g2{quarter} + ?
                             WHERE g1 = ? AND codkon = ?"""

    cur.executemany(sqlite_update_query, record_list_quarter)
    conn.commit()

    sqlite_update_g2_query = f"""UPDATE gr_proj SET g2 = g21 + g22 + g23 + g24
                             WHERE g1 = ? AND codkon = ?"""
    cur.executemany(sqlite_update_g2_query, record_list_g2)
    conn.commit()

    cur.close()
    conn.close()

    query = QSqlQuery(f"""SELECT ifnull(SUM(proj.g2), 0) AS sum_fact
                          FROM gr_konk AS konk
                          LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                          LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
    					  {cond_filter if query_sql_filter_save != {} else ""}""")

    while query.next():
        sum_fact = query.value(0)

    fin_form.factlabel.setText(str(sum_fact) + ' ' + 'руб.')

    fin_form.push_button_print_fin.setEnabled(True)

    calc_konk()


def cancel_fin():
    # отмена финансирования

    conn = sqlite3.connect(db_name)
    cur = conn.cursor()

    list_of_tables = cur.execute("""SELECT name FROM sqlite_master WHERE type='table' AND name='temp_proj'""").fetchall()

    cur.close()
    conn.close()

    if list_of_tables == []:
        pass

    else:

        conn = sqlite3.connect(db_name)
        cur = conn.cursor()

        sqlite_drop_table_query = """DROP TABLE gr_proj"""
        sqlite_replace_table_query = """CREATE TABLE gr_proj AS SELECT * FROM temp_proj"""
        sqlite_drop_temp_table_query = """DROP TABLE temp_proj"""

        cur.execute(sqlite_drop_table_query)
        conn.commit()
        cur.execute(sqlite_replace_table_query)
        conn.commit()
        cur.execute(sqlite_drop_temp_table_query)
        conn.commit()

        cur.close()
        conn.close()

        fin_form.tableView.model().clear()

        query = QSqlQuery(f"""SELECT ifnull(SUM(proj.g2), 0) AS sum_fact
                              FROM gr_konk AS konk
                              LEFT JOIN gr_proj AS proj ON proj.codkon = konk.codkon
                              LEFT JOIN vuz ON proj.codvuz = vuz.codvuz
        					  {cond_filter if query_sql_filter_save != {} else ""}""")

        while query.next():
            sum_fact = query.value(0)

        fin_form.factlabel.setText(str(sum_fact) + ' ' + 'руб.')

        fin_form.push_button_save_fin.setEnabled(False)
        fin_form.push_button_cancel_fin.setEnabled(False)
        fin_form.push_button_print_fin.setEnabled(False)

        calc_konk()


def print_fin():
    # выпуск распроряжения по финансированию
    pass


def open_help_view_window():
    # открывает окно показа справки

    window_list[check_active_window()].close()
    help_view_window.showMaximized()


def open_help_prog_view_window():
    # открывает окно показа справки по программе

    window_list[check_active_window()].close()
    help_prog_view_window.showMaximized()


def close_prog():
    # закрывает программу

    conn = sqlite3.connect(db_name)
    cur = conn.cursor()

    sqlite_drop_temp_table_query = """DROP TABLE IF EXISTS temp_proj"""

    cur.execute(sqlite_drop_temp_table_query)

    conn.commit()
    cur.close()
    conn.close()

    if check_active_window() is not None:
        window_list[check_active_window()].close()


def connect_db():
    # подключает базу данных

    con = QSqlDatabase.addDatabase('QSQLITE')
    con.setDatabaseName(db_name)

    if not con.open():
        print('Не удалось подключиться к базе данных')
        return False
    return con


# проверка подключения к базе данных
if not connect_db():
    sys.exit(-1)
else:
    print('Подключение к базе данных прошло успешно')


app = QApplication([])


# инициализация главного окна
main_window = MainWindow()
main_form = MainForm()
main_form.setupUi(main_window)

# инициализация инф об удачно добавлении
add_ok_window=ADdOkWindow()
add_ok_form=AddOkForm()
add_ok_form.setupUi(add_ok_window)
add_ok_form.pushButton_add_ok.clicked.connect(add_ok)


# инициализация окна показа таблицы конкурсов
konk_view_window = KonkViewWindow()
konk_view_form = KonkViewForm()
konk_view_form.setupUi(konk_view_window)


# инициализация окна показа таблицы вузов
vuz_view_window = VuzViewWindow()
vuz_view_form = VuzViewForm()
vuz_view_form.setupUi(vuz_view_window)


# инициализация окна показа таблицы проектов
proj_view_window = ProjViewWindow()
proj_view_form = ProjViewForm()
proj_view_form.setupUi(proj_view_window)
proj_view_form.push_button_sort.clicked.connect(open_proj_sort_window)
proj_view_form.push_button_filter.clicked.connect(open_proj_filter_window)
proj_view_form.push_button_select_check.clicked.connect(proj_select_check)
proj_view_form.pushButton_add_row.clicked.connect(add_row)


# инициализация окна сортировки таблицы проектов
proj_sort_window = ProjSortWindow()
proj_sort_form = ProjSortForm()
proj_sort_form.setupUi(proj_sort_window)
proj_sort_form.push_button_sort_save.clicked.connect(proj_sort_save)
proj_sort_form.push_button_sort_reset.clicked.connect(proj_sort_reset)


# инициализация окна фильтрации таблицы проектов
proj_filter_window = ProjFilterWindow()
proj_filter_form = ProjFilterForm()
proj_filter_form.setupUi(proj_filter_window)
proj_filter_form.push_button_set_codkon.clicked.connect(proj_filter_codkon)
proj_filter_form.push_button_set_region.clicked.connect(proj_filter_region)
proj_filter_form.push_button_set_oblname.clicked.connect(proj_filter_oblname)
proj_filter_form.push_button_set_city.clicked.connect(proj_filter_city)
proj_filter_form.push_button_set_z2.clicked.connect(proj_filter_z2)
proj_filter_form.push_button_filter_save.clicked.connect(proj_filter_save)
proj_filter_form.push_button_filter_check.clicked.connect(proj_filter_check)
proj_filter_form.push_button_filter_reset.clicked.connect(proj_filter_reset)


# инициализация окна показа справки
help_view_window = HelpViewWindow()
help_view_form = HelpViewForm()
help_view_form.setupUi(help_view_window)


# инициализация окна добавления строки
add_view_window = AddRowWindowWindow()
add_view_form= AddRowWindowForm()
add_view_form.setupUi(add_view_window)
add_view_form.pushButton_add_row_add.clicked.connect(read_data_add_row)
add_view_form.pushButton_back.clicked.connect(return_back)


error_add_row_window=ErrorAddRowWindow()
error_add_row_form=ErrorAddRowForm()
error_add_row_form.setupUi(error_add_row_window)


# инициализация окна показа справки по программе
help_prog_view_window = HelpProgViewWindow()
help_prog_view_form = HelpProgViewForm()
help_prog_view_form.setupUi(help_prog_view_window)


# инициализация редактирования строк
change_window=ChangeWindow()
change_form=ChangeForm()
change_form.setupUi(change_window)
proj_view_form.pushButton_change.clicked.connect(change)
change_form.pushButton_change_back.clicked.connect(return_back1)
change_form.pushButton_change_2.clicked.connect(read_data_add_row1)


# инициализация окна удаления
delete_proj_window=DeleteWindow()
delete_proj_form=DeleteWindowForm()
delete_proj_form.setupUi(delete_proj_window)
delete_proj_form.pushButton_no.clicked.connect(pb_no)
delete_proj_form.pushButton_yes.clicked.connect(pb_yes)


# инициализация окна анализа
analysis_window=AnalysisWindow()
analysis_form=AnalysisForm()
analysis_form.setupUi(analysis_window)



# инициализация окна финансирования
fin_window=FinWindow()
fin_form=FinForm()
fin_form.setupUi(fin_window)
fin_form.push_button_calc_fin.clicked.connect(calc_fin)
fin_form.push_button_save_fin.clicked.connect(save_fin)
fin_form.push_button_cancel_fin.clicked.connect(cancel_fin)
fin_form.push_button_print_fin.clicked.connect(print_fin)


# список всех окон приложения
window_list = [main_window, konk_view_window, vuz_view_window, proj_view_window,  help_view_window, help_prog_view_window, proj_sort_window, proj_filter_window,
               error_add_row_window, add_view_window, change_window, delete_proj_window, analysis_window, fin_window]


# привязка функций показа таблиц ко кнопкам панели в главном окне
main_form.panel_data_konk.triggered.connect(open_konk_window)
main_form.panel_data_vuz.triggered.connect(open_vuz_window)
main_form.panel_data_proj.triggered.connect(open_proj_window)
main_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
main_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
main_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
main_form.panel_finance_action.triggered.connect(open_fin_window)
main_form.panel_help_view.triggered.connect(open_help_view_window)
main_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
main_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне показа таблицы конкурсов
konk_view_form.panel_data_konk.triggered.connect(open_konk_window)
konk_view_form.panel_data_vuz.triggered.connect(open_vuz_window)
konk_view_form.panel_data_proj.triggered.connect(open_proj_window)
konk_view_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
konk_view_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
konk_view_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
konk_view_form.panel_finance_action.triggered.connect(open_fin_window)
konk_view_form.panel_help_view.triggered.connect(open_help_view_window)
konk_view_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
konk_view_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне показа таблицы вузов
vuz_view_form.panel_data_konk.triggered.connect(open_konk_window)
vuz_view_form.panel_data_vuz.triggered.connect(open_vuz_window)
vuz_view_form.panel_data_proj.triggered.connect(open_proj_window)
vuz_view_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
vuz_view_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
vuz_view_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
vuz_view_form.panel_finance_action.triggered.connect(open_fin_window)
vuz_view_form.panel_help_view.triggered.connect(open_help_view_window)
vuz_view_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
vuz_view_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне показа таблицы проектов
proj_view_form.panel_data_konk.triggered.connect(open_konk_window)
proj_view_form.panel_data_vuz.triggered.connect(open_vuz_window)
proj_view_form.panel_data_proj.triggered.connect(open_proj_window)
proj_view_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
proj_view_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
proj_view_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
proj_view_form.panel_finance_action.triggered.connect(open_fin_window)
proj_view_form.panel_help_view.triggered.connect(open_help_view_window)
proj_view_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
proj_view_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне сортировки таблицы проектов
proj_sort_form.panel_data_konk.triggered.connect(open_konk_window)
proj_sort_form.panel_data_vuz.triggered.connect(open_vuz_window)
proj_sort_form.panel_data_proj.triggered.connect(open_proj_window)
proj_sort_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
proj_sort_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
proj_sort_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
proj_sort_form.panel_finance_action.triggered.connect(open_fin_window)
proj_sort_form.panel_help_view.triggered.connect(open_help_view_window)
proj_sort_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
proj_sort_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне фильтрации таблицы проектов
proj_filter_form.panel_data_konk.triggered.connect(open_konk_window)
proj_filter_form.panel_data_vuz.triggered.connect(open_vuz_window)
proj_filter_form.panel_data_proj.triggered.connect(open_proj_window)
proj_filter_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
proj_filter_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
proj_filter_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
proj_filter_form.panel_finance_action.triggered.connect(open_fin_window)
proj_filter_form.panel_help_view.triggered.connect(open_help_view_window)
proj_filter_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
proj_filter_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне анализа
analysis_form.panel_data_konk.triggered.connect(open_konk_window)
analysis_form.panel_data_vuz.triggered.connect(open_vuz_window)
analysis_form.panel_data_proj.triggered.connect(open_proj_window)
analysis_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
analysis_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
analysis_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
analysis_form.panel_finance_action.triggered.connect(open_fin_window)
analysis_form.panel_help_view.triggered.connect(open_help_view_window)
analysis_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
analysis_form.panel_exit_action.triggered.connect(close_prog)


fin_form.panel_data_konk.triggered.connect(open_konk_window)
fin_form.panel_data_vuz.triggered.connect(open_vuz_window)
fin_form.panel_data_proj.triggered.connect(open_proj_window)
fin_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
fin_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
fin_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
fin_form.panel_finance_action.triggered.connect(open_fin_window)
fin_form.panel_help_view.triggered.connect(open_help_view_window)
fin_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
fin_form.panel_exit_action.triggered.connect(close_prog)

# привязка функций показа таблиц ко кнопкам панели в окне показа помощи
help_view_form.panel_data_konk.triggered.connect(open_konk_window)
help_view_form.panel_data_vuz.triggered.connect(open_vuz_window)
help_view_form.panel_data_proj.triggered.connect(open_proj_window)
help_view_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
help_view_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
help_view_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
help_view_form.panel_finance_action.triggered.connect(open_fin_window)
help_view_form.panel_help_view.triggered.connect(open_help_view_window)
help_view_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
help_view_form.panel_exit_action.triggered.connect(close_prog)


# привязка функций показа таблиц ко кнопкам панели в окне показа помощи по программе
help_prog_view_form.panel_data_konk.triggered.connect(open_konk_window)
help_prog_view_form.panel_data_vuz.triggered.connect(open_vuz_window)
help_prog_view_form.panel_data_proj.triggered.connect(open_proj_window)
help_prog_view_form.panel_analysis_vuz.triggered.connect(open_analysis_vuz_window)
help_prog_view_form.panel_analysis_konk.triggered.connect(open_analysis_konk_window)
help_prog_view_form.panel_analysis_subj.triggered.connect(open_analysis_subj_window)
help_prog_view_form.panel_finance_action.triggered.connect(open_fin_window)
help_prog_view_form.panel_help_view.triggered.connect(open_help_view_window)
help_prog_view_form.panel_help_prog.triggered.connect(open_help_prog_view_window)
help_prog_view_form.panel_exit_action.triggered.connect(close_prog)


app.aboutToQuit.connect(close_prog)

open_konk_window()
app.exec()
